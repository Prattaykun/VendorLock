"""
VendorLock Master API & Architecture Documentation Generator.

Generates a single, highly detailed .docx file.
Extracts OpenAPI schemas and maps them deeply to `vendorlock.txt` logic,
frontend components, and backend Agent implementations.
"""
import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

# Deeply detailed Agent Information derived from vendorlock.txt
AGENT_METADATA = {
    "agent1": {
        "name": "Agent 1: Trade Capture & Normalisation (MVP)",
        "vendorlock_section": "Section: 1. Captures All Trade on Chat (MVP) & The 6 Agents — Technical Overview",
        "purpose": "Convert informal, multilingual chat messages (Hindi, Bhojpuri, Tamil, Hinglish) into clean, structured trade records in real-time.",
        "inputs": "Telegram/WhatsApp messages (text, Hinglish, regional languages), voice note transcripts, invoice photos.",
        "logic": "Intent classification (order / payment / return / dispute / scheme query) -> entity extraction (items, quantities, unit prices, retailer ID, payment type, dates) -> validation against product master -> save to ledger.",
        "confirm_commit": "Every parsed order is sent back to the retailer for YES/DISPUTE before final logging.",
        "setup_reqs": "Gemini 2.5 Flash via google-genai, Postgres `orders` table, Telegram Webhook configuration.",
        "triggers": ["/orders/", "/webhook/telegram", "/agent/parse-message"]
    },
    "agent2": {
        "name": "Agent 2: Trust & Behaviour Scoring (MVP)",
        "vendorlock_section": "Section: 2. Builds a CIBIL-Style Trust Score",
        "purpose": "Maintain a CIBIL-style multi-factor Trust Score (0-100) for every counterparty to drive credit decisions and flag risky buyers.",
        "inputs": "Full ledger — payments, orders, returns, disputes, communication logs.",
        "logic": "Computes weighted sub-scores (Payment discipline 30%, Order consistency 20%, Cancellation rate 15%, Return frequency 15%, Communication reliability 10%, Trade stability 10%). Builds composite score and 30/90-day trend. Tiers: A (80-100), B (60-79), C (40-59), D (0-39).",
        "setup_reqs": "Postgres `trust_scores`, `orders`, `returns` tables. Agent 2 LangGraph state graph.",
        "triggers": ["/trust-score/", "recalculate"]
    },
    "agent3": {
        "name": "Agent 3: Risk, Scheme & Compliance Intelligence (Core v2)",
        "vendorlock_section": "Sections: 3 (Scheme Leakage), 5 (Fake Returns), 6 (Expiry Risk)",
        "purpose": "Combine Trust Scores, ledger data, scheme rules and compliance info to surface hidden risks and financial leakages.",
        "inputs": "Trust outputs from Agent 2, full trade ledger, scheme PDFs (via RAG), batch/expiry data, vendor GSTIN.",
        "logic": "1. Credit risk detection. 2. Scheme leakage & pass-through tracking. 3. Returns & damage validation (classifies GENUINE / SUSPICIOUS). 4. Expiry intelligence. 5. GST compliance flag.",
        "setup_reqs": "Vector Store (FAISS/Pinecone) for Scheme PDFs, Postgres `risk_alerts`, `schemes`, `returns`, `batches`.",
        "triggers": ["/risk-alerts/run-scan", "/schemes/", "/returns/", "/expiry/"]
    },
    "agent4": {
        "name": "Agent 4: Action & Recommendation (v2)",
        "vendorlock_section": "Section: Agent 4 — Action & Recommendation",
        "purpose": "Turn risk and intelligence signals from Agent 3 into concrete, human-reviewable actions and draft chat messages.",
        "inputs": "Alerts from Agent 3 + distributor credit policy and risk appetite.",
        "logic": "Generates dashboard action cards and chat-ready draft messages in Hindi/Hinglish (e.g., for collections or return claims). Nothing executes automatically; requires human-in-the-loop approval.",
        "setup_reqs": "Postgres `risk_alerts` table.",
        "triggers": ["/risk-alerts/"]
    },
    "agent5": {
        "name": "Agent 5: Demand & Pre-Stock Forecast (v2)",
        "vendorlock_section": "Section: Agent 5 — Demand & Pre-Stock Forecast",
        "purpose": "Predict future demand per SKU per zone, recommend pre-stocking, and infer secondary sales velocity.",
        "inputs": "Historical order series per SKU-retailer pair, seasonality calendar, active schemes.",
        "logic": "Identifies demand spikes/drops, generates pre-stock alerts, and feeds Agent 2's seasonality baseline.",
        "setup_reqs": "Deep historical order dataset. High-volume read access to `orders`.",
        "triggers": ["/analytics/secondary-sales-estimate", "/agent/run"]
    },
    "agent6": {
        "name": "Agent 6: Beat Intelligence & Coverage (v2)",
        "vendorlock_section": "Section: 4. Monitors Beat Coverage — Catches Ghost Visits",
        "purpose": "Ensure salesmen visit the right outlets, optimise routes, and detect skipped/ghost visits to recover missed revenue.",
        "inputs": "Outlet master list, check-in logs, order history per outlet.",
        "logic": "Detects coverage gaps, cross-references GPS check-ins with 2-way chat traffic to flag ghost visits. Generates daily route plans.",
        "setup_reqs": "Postgres `outlets`, `beat_checkins`, `salesmen` tables.",
        "triggers": ["/beat-plan/"]
    }
}

FRONTEND_UI_MAPPING = {
    "/auth": "Login Screen / Registration Flow (AuthContext in api-client.ts)",
    "/orders": "DistributorControlTower.tsx -> Real-time Orders Ledger Component",
    "/trust-score": "CreditDecisionPanel.tsx -> Trust Radar Chart & MYSCORE Telegram view",
    "/risk-alerts": "DistributorControlTower.tsx -> Risk Alerts Sidebar",
    "/beat-plan": "BeatIntelligencePanel.tsx -> Coverage Heatmap & Salesman Route List",
    "/schemes": "SchemeLeakageMonitor.tsx -> Scheme Pass-through Chart & PDF Upload Modal",
    "/returns": "ExpiryAndReturnsCalendar.tsx -> Returns Approval Workflow Component",
    "/expiry": "ExpiryAndReturnsCalendar.tsx -> Batch Tracking & 90-day Alerts Panel",
    "/certificate": "RetailerProfile.tsx -> Trust Certificate Generator & Public Verification Page",
    "/distributor": "DistributorControlTower.tsx -> KPI Headers, Salesmen Table, Retailer List",
    "/retailer": "RetailerProfile.tsx -> Ledger History View, Credit Limit Editor Modal",
    "/analytics/quick-commerce": "QuickCommerceThreatMap.tsx -> Pin-code Pricing Monitor Visual",
    "/analytics/audit-trail": "AuditTrailViewer.tsx -> Immutable Hash-Chain Log View",
    "/webhook": "External Telegram Bot Interface (No Frontend UI)"
}

def get_agent_for_path(path):
    agents = []
    for ag_id, data in AGENT_METADATA.items():
        for trigger in data["triggers"]:
            if trigger in path:
                agents.append(data)
                break
    return agents

def get_ui_for_path(path):
    for key, ui in FRONTEND_UI_MAPPING.items():
        if path.startswith(key):
            return ui
    return "Generic Dashboard View"

def resolve_ref(ref_str, openapi):
    parts = ref_str.split('/')
    if len(parts) == 4 and parts[1] == 'components' and parts[2] == 'schemas':
        return openapi.get('components', {}).get('schemas', {}).get(parts[3], {})
    return {}

def extract_schema_fields(schema, openapi):
    """Flattens OpenAPI schema into a list of (Field, Type, Required, Description)."""
    if '$ref' in schema:
        schema = resolve_ref(schema['$ref'], openapi)
        
    fields = []
    if schema.get('type') == 'object' and 'properties' in schema:
        req_fields = schema.get('required', [])
        for k, v in schema['properties'].items():
            field_type = v.get('type', 'any')
            if '$ref' in v:
                field_type = v['$ref'].split('/')[-1]
            elif field_type == 'array' and 'items' in v:
                if '$ref' in v['items']:
                    field_type = f"Array<{v['items']['$ref'].split('/')[-1]}>"
                else:
                    field_type = f"Array<{v['items'].get('type', 'any')}>"
            
            desc = v.get('title', '') or v.get('description', '')
            req_str = "Yes" if k in req_fields else "No"
            fields.append((k, field_type, req_str, desc))
    return fields

def add_schema_table(doc, title, fields):
    doc.add_heading(title, level=3)
    if not fields:
        p = doc.add_paragraph("No specific body fields (or unstructured JSON/FormData).")
        p.runs[0].italic = True
        return
        
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Required'
    hdr_cells[3].text = 'Description'
    
    for f in fields:
        row_cells = table.add_row().cells
        row_cells[0].text = f[0]
        row_cells[1].text = f[1]
        row_cells[2].text = f[2]
        row_cells[3].text = str(f[3])

def main():
    os.environ.setdefault("APP_ENV", "development")
    print("[*] Loading FastAPI app to extract OpenAPI schema...")
    try:
        from main import app
        openapi = app.openapi()
    except Exception as e:
        print(f"[ERROR] Failed to load app: {e}")
        return

    doc = Document()
    
    # Title Page
    title = doc.add_heading('VendorLock: Master API & Architecture Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Generated automatically from FastAPI Introspection and vendorlock.txt Spec", style='Subtitle')
    
    doc.add_page_break()
    
    # Exec Summary
    doc.add_heading('1. Executive Architecture Summary', level=1)
    doc.add_paragraph(
        "VendorLock is a chat-native AI intelligence layer for FMCG and HORECA. "
        "It acts as a multi-tenant SaaS that ingests unstructured Telegram/WhatsApp trade messages, "
        "converts them into a structured ledger via Agent 1, and subsequently triggers Agents 2-6 "
        "to calculate Trust Scores, detect scheme leakage, flag expiry risks, and monitor route coverage."
    )
    
    doc.add_heading('Core Infrastructure Setup', level=2)
    infra_table = doc.add_table(rows=5, cols=2)
    infra_table.style = 'Table Grid'
    infra_data = [
        ("Framework", "FastAPI (Python 3.11+), Pydantic V2"),
        ("Primary LLM", "Gemini 2.5 Flash via Google GenAI SDK (Express Mode)"),
        ("Fallback LLM", "NVIDIA NIM (DeepSeek V4 Pro) via OpenAI-compatible SDK"),
        ("Relational DB", "Supabase (PostgreSQL) - Ledger, Auth, Alerts"),
        ("NoSQL DB", "MongoDB Atlas - Unstructured chat payloads, PDF raw text")
    ]
    for i, (k, v) in enumerate(infra_data):
        infra_table.rows[i].cells[0].text = k
        infra_table.rows[i].cells[1].text = v
        
    doc.add_page_break()
    
    # Endpoints grouped by tags
    paths = openapi.get("paths", {})
    tags_grouped = {}
    for path, methods in paths.items():
        for method, details in methods.items():
            tag = details.get("tags", ["Uncategorized"])[0]
            if tag not in tags_grouped:
                tags_grouped[tag] = []
            details['path'] = path
            details['method'] = method.upper()
            tags_grouped[tag].append(details)

    doc.add_heading('2. API Endpoints by Module', level=1)
    
    for tag, endpoints in tags_grouped.items():
        doc.add_heading(f"{tag.upper()} Module", level=2)
        
        for ep in endpoints:
            # Header
            p = doc.add_paragraph()
            p.add_run(f"[{ep['method']}] ").bold = True
            p.add_run(ep['path']).font.color.rgb = RGBColor(0, 0, 255)
            
            doc.add_paragraph(ep.get('summary', 'No summary provided.'), style='Intense Quote')
            if ep.get('description'):
                doc.add_paragraph(ep['description'])
                
            # Role & Frontend
            doc.add_heading("Context & Roles", level=3)
            fe = get_ui_for_path(ep['path'])
            doc.add_paragraph(f"Frontend Component: ", style='List Bullet').runs[0].bold = True
            doc.paragraphs[-1].add_run(fe)
            
            # Agents
            agents = get_agent_for_path(ep['path'])
            if agents:
                doc.add_heading("AI Agent Orchestration", level=3)
                for ag in agents:
                    doc.add_paragraph(f"Driven by {ag['name']}", style='List Bullet').runs[0].bold = True
                    doc.add_paragraph(f"Spec Reference: {ag['vendorlock_section']}", style='List Bullet 2')
                    doc.add_paragraph(f"Logic: {ag['logic']}", style='List Bullet 2')
                    doc.add_paragraph(f"Required Setup: {ag['setup_reqs']}", style='List Bullet 2')
            
            # Request Body
            req_body = ep.get('requestBody', {})
            if req_body:
                content = req_body.get('content', {})
                if 'application/json' in content:
                    schema = content['application/json'].get('schema', {})
                    fields = extract_schema_fields(schema, openapi)
                    add_schema_table(doc, "Request JSON Schema", fields)
                elif 'multipart/form-data' in content:
                    doc.add_heading("Request Body", level=3)
                    doc.add_paragraph("multipart/form-data (Requires file upload binary).")
                    
            # Response Body
            responses = ep.get('responses', {})
            success_resp = responses.get("200") or responses.get("201", {})
            if success_resp:
                content = success_resp.get('content', {})
                if 'application/json' in content:
                    schema = content['application/json'].get('schema', {})
                    fields = extract_schema_fields(schema, openapi)
                    add_schema_table(doc, "Response JSON Schema", fields)
                    
            doc.add_paragraph("-" * 60)
            
    out_dir = ROOT / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = out_dir / "VendorLock_API_Master_Documentation.docx"
    doc.save(str(filename))
    print(f"\n[OK] Highly detailed DOCX generated at: {filename}")

if __name__ == "__main__":
    main()
